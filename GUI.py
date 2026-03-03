from tkinter import filedialog
from tkinter import messagebox

from pathlib import Path

from google import genai

from settings.base import MainSettings

from dotenv import load_dotenv 
import os

import threading

load_dotenv(Path(__file__).parent / "settings" / "tokens.env")

API_KEY = os.getenv("GEMINI")

client = genai.Client(api_key=API_KEY)



class GUI(MainSettings):
  def __init__(self,x,y,name):
    super().__init__(x,y,name)

  def uploadBtn(self):
    full_info = {}
    try:
      choice_type = self.type_var.get()
      choice_task = self.task_var.get()

      if choice_type and choice_task:
        match choice_type:
          case 1:
            waiting = [("Word","*.docx")]
          case 2:
            waiting = [("Text","*.txt")]
          case 3:
            waiting = [("PDF","*.pdf")]

        filename = filedialog.askopenfile(title="Choose file",filetypes=waiting)

        if filename:
          try:
            if choice_type == 1:
              info = self.read_docx(filename)

            if choice_type == 2:
              info = filename.read()

            if choice_type == 3:
              from pypdf import PdfReader
              reader = PdfReader(filename.name)
              info = ""

              for page in reader.pages:
                info += page.extract_text() + "\n"
          except Exception as e:
            print(f"❌ [ERROR] Getting info {e}")
            return
          
          
        match choice_task:
          case 1:
            full_info["task"] = 1
            full_info["info"] = info
          case 2:
            full_info["task"] = 2
            full_info["info"] = info
          case 3:
            full_info["task"] = 3
            full_info["info"] = info

        
        if info:
          self.ai_check(full_info)


      else:
          if not choice_type:
            messagebox.showerror("❌ [ERROR]","You must choose the type of file for upload!")
          if not choice_task:
            messagebox.showerror("❌ [ERROR]","You must choose the task for upload!")
            return
          
    except Exception as e:
      print("❌ [ERROR] Choosing buttons")


  def read_docx(self, filename):
    import docx
    doc = docx.Document(filename)
    full_text = []

    for para in doc.paragraphs:
      full_text.append(para.text)


    for table in doc.tables:
      for row in table.rows:
        for cell in row.cells:
          full_text.append(cell.text)

    return '\n'.join(full_text)

  def ai_check(self,info):
    task = info["task"]
    info = info["info"]

    self.progress.pack(pady=10)
    self.progress.start(10)

    def run():
      response = None

      try:
        match task:
          case 1:

            task_file = client.files.upload(file="prompts/Rating+Advice.txt")

            response = client.models.generate_content(
              model = "gemini-2.5-flash-lite",
              contents = [task_file,
                        f"{info}, Your task is to complete the assignment described in the file, based on the business information I sent you."]
            )

            with open("response.txt", "w", encoding="utf-8") as file:
              file.write(response.text)
          case 2:
            task_file = client.files.upload(file="prompts/site.txt")

            response = client.models.generate_content(
              model = "gemini-2.5-flash-lite",
              contents = [task_file,
                        f"{info}\n\n, Your task is to complete the assignment described in the file, based on the business information I sent you."]
            )

            with open("response.txt", "w", encoding="utf-8") as file:
              file.write(response.text)
          case 3:

            task_file = client.files.upload(file="prompts/PESTEL AND SWOT.txt")

            response = client.models.generate_content(
              model = "gemini-2.5-flash-lite",
              contents = [task_file,
                        f"{info}\n\n, Your task is to complete the assignment described in the file, based on the business information I sent you."]
            )

            with open("response.txt", "w", encoding="utf-8") as file:
              file.write(response.text)

      except Exception as e:
        print(f"❌ [ERROR] AI GENERATE {e}")
        return None
      finally:
        self.progress.after(0, self.progress.stop)
        self.progress.after(0, self.progress.pack_forget)
        self.progress.after(0, self.getResult)

    thread = threading.Thread(target=run, daemon=True)
    thread.start()
      
  
  def getResult(self):
    src = Path("response.txt")
    if not src.exists():
      messagebox.showerror("⚠️", "No result yet!")
      return
    
    try:
      with open("response.txt", "r", encoding="utf-8") as file:
        response = file.read()
    except FileNotFoundError:
        messagebox.showerror("⚠️", "PROBLEM WITH AI RESPONSE")
        return
    
  
    dest = filedialog.asksaveasfilename(
      title="Save result",
      defaultextension=".txt",
      filetypes=[("Text", "*.txt")]
    )

    if dest:
      import shutil
      shutil.copy(src, dest)
      messagebox.showinfo("✅", "File saved!")


  def checkExamples(self):
    import webbrowser
    webbrowser.open("https://drive.google.com/drive/folders/1odia6HumiOPGQwh-zCghA9H6Q1mmxKfl?usp=drive_link")

    

gui = GUI(800,550,"Program")

gui.start()